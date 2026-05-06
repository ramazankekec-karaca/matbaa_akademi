import os
import uuid
from django.conf import settings
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.utils import timezone

def _create_base_document(sinav_bilgileri):
    """Tüm sınav formatları için ortak olan üst başlık (Header) kısmını oluşturur."""
    doc = Document()
    
    # Sayfa kenar boşluklarını ayarla
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Üst Başlık Bilgileri
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    okul_adi = sinav_bilgileri.get('okulAdi', 'OKUL ADI GİRİLMEDİ').upper()
    egitim_yili = sinav_bilgileri.get('egitimYili', '202X - 202X').upper()
    bolum_alan = sinav_bilgileri.get('bolumAlan', 'MATBAA TEKNOLOJİSİ ALANI').upper()
    ders_adi = sinav_bilgileri.get('dersAdi', 'DERS ADI').upper()
    sinav_donemi = sinav_bilgileri.get('sinavDonemi', '1. DÖNEM 1. YAZILI').upper()

    run = header_para.add_run(f"{okul_adi}\n")
    run.bold = True
    run.font.size = Pt(12)
    
    run = header_para.add_run(f"{egitim_yili} EĞİTİM ÖĞRETİM YILI\n")
    run.bold = True
    
    run = header_para.add_run(f"{bolum_alan}\n")
    run.bold = True
    
    run = header_para.add_run(f"{ders_adi} DERSİ {sinav_donemi} YOKLAMA KAĞIDI\n")
    run.bold = True

    # Öğrenci Bilgi Alanı (Tablo Şeklinde)
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    
    row0 = table.rows[0].cells
    row0[0].text = "ADI SOYADI:"
    row0[1].text = "SINIFI / NO:"
    row0[2].text = f"SÜRE: {sinav_bilgileri.get('sinavSuresi', '40')} DAKİKA"
    
    row1 = table.rows[1].cells
    row1[0].text = "ALDIĞI PUAN:"
    row1[1].text = "NOTLA DEĞERİ:"
    row1[2].text = "ÖĞRETMEN İMZA:"

    doc.add_paragraph("\nSORULAR")
    return doc

def _add_questions_to_doc(doc, secilen_sorular):
    """Soruları Word belgesine ekler."""
    for i, soru_data in enumerate(secilen_sorular, 1):
        # Soru Metni
        puan = soru_data.get('puan', 0)
        soru_metni = f"Soru {i} ({puan} Puan): {soru_data.get('soru_metni', '')}"
        
        p = doc.add_paragraph()
        run = p.add_run(soru_metni)
        run.bold = True
        
        # Soru Resmi Varsa Ekle
        resim_yolu = soru_data.get('soru_resmi')
        if resim_yolu:
            # Resim yolunu fiziksel yola çevir (media/ klasöründen)
            full_path = os.path.join(settings.MEDIA_ROOT, resim_yolu.replace('/media/', ''))
            if os.path.exists(full_path):
                doc.add_picture(full_path, width=Cm(6))

        # Şıklar/Cevaplar
        cevaplar = soru_data.get('cevaplar', [])
        for j, cevap in enumerate(cevaplar):
            harf = chr(65 + j) # A, B, C...
            cevap_metni = f"{harf}) {cevap.get('metin', '')}"
            cp = doc.add_paragraph(cevap_metni)
            cp.paragraph_format.left_indent = Cm(1)
            
        doc.add_paragraph("") # Sorular arası boşluk

def _save_document(doc, prefix):
    """Belgeyi media/sinavlar dizinine kaydeder ve URL döner."""
    save_dir = os.path.join(settings.MEDIA_ROOT, 'sinavlar')
    os.makedirs(save_dir, exist_ok=True)
    
    filename = f"{prefix}_{timezone.now().strftime('%Y%m%d')}_{uuid.uuid4().hex[:6]}.docx"
    file_path = os.path.join(save_dir, filename)
    doc.save(file_path)
    
    return f"{settings.MEDIA_URL}sinavlar/{filename}"

# --- VIEWS.PY TARAFINDAN ÇAĞRILAN ANA FONKSİYONLAR ---

def sinav_olustur(sinav_bilgileri, secilen_sorular):
    doc = _create_base_document(sinav_bilgileri)
    _add_questions_to_doc(doc, secilen_sorular)
    return _save_document(doc, "Standart_Sinav")

def sinav_olustur2(sinav_bilgileri, secilen_sorular):
    # Alternatif standart sınav formatı (Aynı mantığı kullanabilirsiniz veya özelleştirebilirsiniz)
    return sinav_olustur(sinav_bilgileri, secilen_sorular)

def sorumluluk_sinavi_olustur(sinav_bilgileri, secilen_sorular):
    # Sorumluluk sınavı için başlıkları özelleştir
    sinav_bilgileri['sinavDonemi'] = "SORUMLULUK SINAVI"
    doc = _create_base_document(sinav_bilgileri)
    
    komisyon_para = doc.add_paragraph()
    komisyon_para.add_run(f"\nKomisyon Başkanı: {sinav_bilgileri.get('komisyonBaskani', '')}\n").bold = True
    komisyon_para.add_run(f"Komisyon Üyeleri:\n{sinav_bilgileri.get('komisyonUyeleri', '')}")
    
    _add_questions_to_doc(doc, secilen_sorular)
    return _save_document(doc, "Sorumluluk_Sinavi")

def notbaremi_sinavi_olustur(sinav_bilgileri, secilen_sorular):
    doc = _create_base_document(sinav_bilgileri)
    _add_questions_to_doc(doc, secilen_sorular)
    
    doc.add_page_break()
    doc.add_heading('Cevap Anahtarı ve Not Baremi', level=1)
    
    # Not baremi tablosu
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Soru No'
    hdr_cells[1].text = 'Doğru Cevap'
    hdr_cells[2].text = 'Puan'
    
    for i, soru in enumerate(secilen_sorular, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = soru.get('dogru_cevap_harf', 'Klasik/Bilinmiyor')
        row_cells[2].text = str(soru.get('puan', 0))
        
    return _save_document(doc, "Not_Baremli_Sinav")

def calisma_sorulari_olustur(sinav_bilgileri, secilen_sorular):
    doc = _create_base_document(sinav_bilgileri)
    doc.paragraphs[0].clear()
    
    baslik = doc.add_paragraph()
    baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = baslik.add_run(f"{sinav_bilgileri.get('dersAdi', 'DERS')} ÇALIŞMA SORULARI\n")
    run.bold = True
    run.font.size = Pt(14)

    # Sorular ve hemen altına doğru cevabı ekle
    for i, soru_data in enumerate(secilen_sorular, 1):
        soru_metni = f"Soru {i}: {soru_data.get('soru_metni', '')}"
        p = doc.add_paragraph()
        p.add_run(soru_metni).bold = True
        
        cevaplar = soru_data.get('cevaplar', [])
        for j, cevap in enumerate(cevaplar):
            harf = chr(65 + j)
            doc.add_paragraph(f"{harf}) {cevap.get('metin', '')}").paragraph_format.left_indent = Cm(1)
            
        dogru = soru_data.get('dogru_cevap_harf', '')
        if dogru:
            doc.add_paragraph(f"Cevap: {dogru}").bold = True
            
        doc.add_paragraph("")
        
    return _save_document(doc, "Calisma_Sorulari")