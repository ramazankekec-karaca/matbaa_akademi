import io
import os
import re
import json
from docx import Document
from docx.shared import Inches
from django.conf import settings
from django.http import HttpResponse, JsonResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth import login, authenticate, logout
from django.core.files.storage import FileSystemStorage
from .models import Soru, Sinif, Dal, Ders, Modul, CustomUser, SoruTuru, HataBildirimi
from .forms import CustomUserCreationForm, SoruForm, ContactForm
from django.contrib import messages

# 1. Ana Sayfa ve Temel Linkler
def home(request):
    """Ana sayfayı yükler ve ders istatistiklerini getirir."""
    # select_related ile veritabanını yormadan dal ve sınıf bilgilerini de tek seferde çekiyoruz
    dersler = Ders.objects.select_related('dal__sinif').all()
    ders_istatistikleri = []
    
    for ders in dersler:
        soru_adet = Soru.objects.filter(ders=ders, onay_durumu='onaylandi').count()
        ders_istatistikleri.append({
            'adi': ders.adi,
            'soru_sayisi': soru_adet,
            'dal_adi': ders.dal.adi,          # YENİ: Dal adını aldık
            'sinif': ders.dal.sinif.sinif     # YENİ: Sınıf derecesini aldık
        })
        
    context = {
        'ders_istatistikleri': ders_istatistikleri
    }
    
    return render(request, 'home.html', context)

def about(request):
    """Hakkımızda sayfasını yükler."""
    return render(request, 'about.html')

def contact_view(request):
    """İletişim formunu işler."""
    if request.method == 'POST':
        form = ContactForm(request.POST)
        if form.is_valid():
            return render(request, 'contact.html', {'success': True})
    else:
        form = ContactForm()
    return render(request, 'contact.html', {'form': form})

@login_required(login_url='login')
def panel(request):
    """Kullanıcı özel panelini yükler."""
    return render(request, 'panel.html')


# 2. Soru Bankası ve Sınav Sihirbazı
def soru_bankasi(request):
    sorular = Soru.objects.filter(onay_durumu='onaylandi').order_by('-id')
    siniflar = Sinif.objects.all()
    soru_turleri = SoruTuru.objects.all()
    
    exam_questions = request.session.get('exam_questions', [])
    
    all_dersler = Ders.objects.all()
    all_dersler_json = json.dumps([
        {
            'id': d.id, 
            'ders_adi': d.adi, 
            'sinif': d.dal.sinif.sinif, 
            'dal_adi': d.dal.adi
        } for d in all_dersler
    ])

    context = {
        'sorular': sorular,
        'siniflar': siniflar,
        'soru_turleri': soru_turleri,
        'all_dersler_json': all_dersler_json,
        'exam_questions': exam_questions,
    }
    return render(request, 'sorubankasi.html', context)

def sinav_sihirbazi(request):
    """Sınav oluşturma sihirbazı sayfasını yükler ve form gönderildiğinde sınavı oluşturur."""
    
    if request.GET.get('soru_sayisi'):
        sinif_id = request.GET.get('sinif')
        dal_id = request.GET.get('dal')
        ders_id = request.GET.get('ders')
        
        # GÜNCELLEME: get() yerine getlist() kullanarak seçilen TÜM modülleri bir liste olarak alıyoruz
        modul_ids = request.GET.getlist('modul') 
        
        tur_id = request.GET.get('tur')
        zorluk = request.GET.get('zorluk')
        
        try:
            soru_sayisi = int(request.GET.get('soru_sayisi', 20))
        except ValueError:
            soru_sayisi = 20

        sorular = Soru.objects.filter(onay_durumu='onaylandi')
        
        if sinif_id:
            sorular = sorular.filter(sinif_id=sinif_id)
        if dal_id:
            sorular = sorular.filter(dal_id=dal_id)
        if ders_id:
            sorular = sorular.filter(ders_id=ders_id)
            
        # GÜNCELLEME: Eğer kullanıcı birden fazla modül seçtiyse, soruları sadece o modüllerin içinden ( __in ) süz
        if modul_ids and len(modul_ids) > 0 and modul_ids[0] != '':
            sorular = sorular.filter(modul_id__in=modul_ids)
            
        if tur_id and tur_id != 'all':
            sorular = sorular.filter(tur_id=tur_id)
        if zorluk and zorluk != 'karisik':
            sorular = sorular.filter(zorluk_seviyesi=zorluk)

        # Soruları karıştırıp istenilen sayıda getir
        secilen_sorular = sorular.order_by('?')[:soru_sayisi]

        return render(request, 'sinav_kagit.html', {
            'sorular': secilen_sorular,
            'manuel_sinav': False 
        })

    siniflar = Sinif.objects.all()
    soru_turleri = SoruTuru.objects.all()
    return render(request, 'sihirbaz.html', {
        'siniflar': siniflar,
        'soru_turleri': soru_turleri
    })

# 3. Soru Ekleme ve Düzenleme İşlemleri
@login_required(login_url='login')
def yeni_soru(request):
    """Tekli veya Metin Kutusundan Toplu Kopyala-Yapıştır ile soru ekler."""
    if request.method == 'POST':
        # 1. Ortak Verileri Al (Hem tekli hem toplu ekleme için geçerli)
        sinif_id = request.POST.get('sinif')
        dal_id = request.POST.get('dal')
        ders_id = request.POST.get('ders')
        modul_id = request.POST.get('modul')
        tur_id = request.POST.get('tur')
        zorluk_seviyesi = request.POST.get('zorluk_seviyesi', 'orta')

        sinif = get_object_or_404(Sinif, id=sinif_id)
        dal = Dal.objects.filter(id=dal_id).first() if dal_id else None  # Alan ortak ise boş olabilir
        ders = get_object_or_404(Ders, id=ders_id)
        modul = get_object_or_404(Modul, id=modul_id)
        tur = get_object_or_404(SoruTuru, id=tur_id)

        # 2. HANGİ MOD ÇALIŞIYOR KONTROL ET (Toplu mu Tekli mi?)
        is_bulk_mode = request.POST.get('bulk_mode') == '1' or request.POST.get('analyze_bulk') == '1'

        if is_bulk_mode:
            # === TOPLU EKLEME MANTIĞI ===
            bulk_text = request.POST.get('bulk_text', '')
            bulk_answer_key = request.POST.get('bulk_answer_key', '')

            if not bulk_text:
                messages.error(request, "Toplu soru metni alanı boş bırakılamaz.")
                return redirect('yeni_soru')

            # Soruları sayı numaralarına göre parçala (Örn: "1.", "2)", "1-")
            soru_bloklari = re.split(r'\n\d+[\.\)\-]\s*', '\n' + bulk_text)
            soru_bloklari = [b.strip() for b in soru_bloklari if b.strip()]

            # Cevap anahtarından sadece harfleri (A,B,C,D,E) ayıkla (Örn: "1-B, 2-A" -> ['B', 'A'])
            cevaplar = re.findall(r'[A-Ea-e]', bulk_answer_key.upper())

            eklenen_sayi = 0
            for i, blok in enumerate(soru_bloklari):
                # Her bir sorunun içindeki şıkları parçala (Örn: "A)", "B.", "C-")
                parts = re.split(r'\n[A-Ea-e][\.\)\-]\s*', '\n' + blok)
                
                soru_metni = parts[0].strip()
                secenekler = [p.strip() for p in parts[1:]] if len(parts) > 1 else []

                # Doğru cevabı sırasıyla eşleştir (Eğer cevap anahtarı girildiyse)
                dogru_cevap = cevaplar[i] if i < len(cevaplar) else None

                Soru.objects.create(
                    sinif=sinif, dal=dal, ders=ders, modul=modul, tur=tur,
                    soru_metni=soru_metni,
                    secenek_a=secenekler[0] if len(secenekler) > 0 else None,
                    secenek_b=secenekler[1] if len(secenekler) > 1 else None,
                    secenek_c=secenekler[2] if len(secenekler) > 2 else None,
                    secenek_d=secenekler[3] if len(secenekler) > 3 else None,
                    secenek_e=secenekler[4] if len(secenekler) > 4 else None,
                    dogru_cevap=dogru_cevap,
                    zorluk_seviyesi=zorluk_seviyesi,
                    soruyu_hazirlayan=request.user,
                    onay_durumu='onaylandi'
                )
                eklenen_sayi += 1

            messages.success(request, f"Harika! {eklenen_sayi} adet soru başarıyla analiz edilip sisteme eklendi.")
            return redirect('soru_bankasi')

        else:
            # === TEKLİ (MANUEL) EKLEME MANTIĞI ===
            Soru.objects.create(
                sinif=sinif, dal=dal, ders=ders, modul=modul, tur=tur,
                soru_metni=request.POST.get('soru_metni'),
                soru_resmi=request.FILES.get('soru_resmi'),
                cevap_resmi=request.FILES.get('cevap_resmi'),
                klasik_cevap=request.POST.get('klasik_cevap'),
                secenek_a=request.POST.get('secenek_a'),
                secenek_b=request.POST.get('secenek_b'),
                secenek_c=request.POST.get('secenek_c'),
                secenek_d=request.POST.get('secenek_d'),
                secenek_e=request.POST.get('secenek_e', ''),
                dogru_cevap=request.POST.get('dogru_cevap'),
                zorluk_seviyesi=zorluk_seviyesi,
                soruyu_hazirlayan=request.user,
                onay_durumu='onaylandi'
            )
            messages.success(request, "Soru başarıyla eklendi.")
            return redirect('soru_bankasi')

    context = {
        'siniflar': Sinif.objects.all(),
        'soru_turleri': SoruTuru.objects.all(),
        'show_ads': False
    }
    return render(request, 'yeni_soru.html', context)

@login_required(login_url='login')
def soru_duzenle(request, soru_id):
    """Mevcut bir soruyu düzenleme formunu işler."""
    soru = get_object_or_404(Soru, id=soru_id)

    if request.method == 'POST':
        soru.sinif_id = request.POST.get('sinif')
        soru.dal_id = request.POST.get('dal')
        soru.ders_id = request.POST.get('ders')
        soru.modul_id = request.POST.get('modul')
        soru.tur_id = request.POST.get('tur')
        soru.zorluk_seviyesi = request.POST.get('zorluk_seviyesi')
        
        soru.soru_metni = request.POST.get('soru_metni')
        soru.klasik_cevap = request.POST.get('klasik_cevap')
        soru.secenek_a = request.POST.get('secenek_a')
        soru.secenek_b = request.POST.get('secenek_b')
        soru.secenek_c = request.POST.get('secenek_c')
        soru.secenek_d = request.POST.get('secenek_d')
        soru.secenek_e = request.POST.get('secenek_e', '') 
        soru.dogru_cevap = request.POST.get('dogru_cevap')

        if request.FILES.get('soru_resmi'):
            soru.soru_resmi = request.FILES.get('soru_resmi')
        if request.FILES.get('cevap_resmi'):
            soru.cevap_resmi = request.FILES.get('cevap_resmi')

        soru.save()
        return redirect('soru_bankasi')

    context = {
        'soru': soru,
        'siniflar': Sinif.objects.all(),
        'soru_turleri': SoruTuru.objects.all(),
    }
    return render(request, 'soru_duzenle.html', context)

@login_required(login_url='login')
def yeni_soru_manual(request):
    """Manuel soru ekleme işlemlerini yönetir."""
    if request.method == 'POST':
        form = SoruForm(request.POST, request.FILES)
        if form.is_valid():
            soru = form.save(commit=False)
            soru.soruyu_hazirlayan = request.user
            soru.save()
            return redirect('soru_bankasi')
    else:
        form = SoruForm()
    return render(request, 'yeni_soru_manual.html', {'form': form})

@login_required(login_url='login')
def yeni_soru_bulk(request):
    """Toplu soru ekleme (Excel vb.) işlemlerini yönetir."""
    if request.method == 'POST':
        return redirect('soru_bankasi')
    return render(request, 'yeni_soru_bulk.html')


# 4. Üyelik İşlemleri
def register(request):
    """Yeni kullanıcı kayıt işlemini gerçekleştirir."""
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect('home')
    else:
        form = CustomUserCreationForm()
    return render(request, 'register.html', {'form': form})


# 5. API ve AJAX (Dinamik Filtreleme)
def get_filtered_sorular(request):
    """Filtreleme isteklerine JSON yanıt döner."""
    sinif_adi = request.GET.get('sinif')
    dal_adi = request.GET.get('dal')
    ders_adi = request.GET.get('ders')
    tur_adi = request.GET.get('tur')
    modul_adlari = request.GET.get('modul', '').split(',') if request.GET.get('modul') else []

    sorular = Soru.objects.filter(onay_durumu='onaylandi')

    if sinif_adi: sorular = sorular.filter(sinif__sinif=sinif_adi)
    if dal_adi: sorular = sorular.filter(dal__adi=dal_adi)
    if ders_adi: sorular = sorular.filter(ders__adi=ders_adi)
    if tur_adi: sorular = sorular.filter(tur__adi=tur_adi)
    if modul_adlari and modul_adlari[0] != '':
        sorular = sorular.filter(modul__adi__in=modul_adlari)

    soru_listesi = []
    for s in sorular:
        soru_listesi.append({
            'id': s.id,
            'soru_metni': s.soru_metni,
            'soru_resmi': s.soru_resmi.url if s.soru_resmi else None,
            'sinif': s.sinif.sinif,
            'dal': s.dal.adi,
            'ders': s.ders.adi,
            'modul': s.modul.adi,
            'tur': s.tur.adi,
            'zorluk_seviyesi': s.get_zorluk_seviyesi_display(),
            'secenek_a': s.secenek_a,
            'secenek_b': s.secenek_b,
            'secenek_c': s.secenek_c,
            'secenek_d': s.secenek_d,
            'secenek_e': s.secenek_e,
            'dogru_cevap': s.dogru_cevap,
            'soruyu_hazirlayan': s.soruyu_hazirlayan.username if s.soruyu_hazirlayan else "Anonim"
        })
    return JsonResponse({'success': True, 'sorular': soru_listesi})


# 6. API ve AJAX (Açılır Menü Doldurma)
def get_siniflar(request): return JsonResponse({'siniflar': list(Sinif.objects.values('id', 'sinif'))})
def get_dallar(request, sinif_id): return JsonResponse({'dallar': list(Dal.objects.filter(sinif_id=sinif_id).values('id', 'adi'))})
def get_dersler(request, dal_id): return JsonResponse({'dersler': list(Ders.objects.filter(dal_id=dal_id).values('id', 'adi'))})
def get_moduller(request, ders_id): return JsonResponse({'moduller': list(Modul.objects.filter(ders_id=ders_id).values('id', 'adi'))})


# 7. Resim Yükleme API'si
@login_required(login_url='login')
def upload_answer_image(request):
    """Sorulara/şıklar dinamik resim ekleme isteklerini karşılar."""
    if request.method == 'POST' and request.FILES.get('image'):
        image = request.FILES['image']
        fs = FileSystemStorage()
        filename = fs.save(image.name, image)
        return JsonResponse({'success': True, 'url': fs.url(filename)})
    return JsonResponse({'success': False, 'error': 'Geçersiz İstek'})


# --- MANUEL SINAV OLUŞTURMA FONKSİYONLARI ---
def toggle_exam_question(request):
    """Sınav sepetine AJAX ile soru ekler veya çıkarır."""
    if request.method == 'POST':
        soru_id = str(request.POST.get('soru_id'))
        if 'exam_questions' not in request.session:
            request.session['exam_questions'] = []
            
        exam_questions = request.session['exam_questions']
        if soru_id in exam_questions:
            exam_questions.remove(soru_id)
            added = False
        else:
            exam_questions.append(soru_id)
            added = True
            
        request.session.modified = True
        return JsonResponse({'success': True, 'added': added, 'count': len(exam_questions)})
    return JsonResponse({'success': False})

def manuel_sinav_kagit(request):
    """Sepetteki sorulardan sınav kağıdı oluşturur."""
    exam_questions = request.session.get('exam_questions', [])
    if not exam_questions:
        return redirect('soru_bankasi')
        
    sorular = Soru.objects.filter(id__in=exam_questions)
    return render(request, 'sinav_kagit.html', {'sorular': sorular, 'manuel_sinav': True})

def clear_exam_session(request):
    """Sepeti tamamen boşaltır."""
    if 'exam_questions' in request.session:
        del request.session['exam_questions']
    return redirect('soru_bankasi')


# --- WORD'E AKTAR (İNDİR) FONKSİYONU ---
def export_word(request):
    """Ekranda görünen soruları Word dosyası olarak indirir. Boş şıkları atlar, resimli cevap anahtarı destekler."""
    if request.method == 'POST':
        soru_ids = request.POST.getlist('soru_ids')
        sorular = Soru.objects.filter(id__in=soru_ids)
    else:
        sinif_adi = request.GET.get('sinif')
        dal_adi = request.GET.get('dal')
        ders_adi = request.GET.get('ders')
        sorular = Soru.objects.filter(onay_durumu='onaylandi').order_by('-id')
        if sinif_adi: sorular = sorular.filter(sinif__sinif=sinif_adi)
        if dal_adi: sorular = sorular.filter(dal__adi=dal_adi)
        if ders_adi: sorular = sorular.filter(ders__adi=ders_adi)

    sablon_yolu = os.path.join(settings.BASE_DIR, 'akademi', 'sablonlar', 'sinav_sablonu.docx')
    
    try:
        document = Document(sablon_yolu)
    except Exception as e:
        document = Document()
        baslik = document.add_paragraph()
        baslik.add_run('Matbaa Akademi - Sınav Kağıdı').bold = True
        document.add_paragraph()

    if not sorular.exists():
        document.add_paragraph('Sınav kağıdına aktarılacak soru bulunamadı.')
    else:
        for i, soru in enumerate(sorular, 1):
            p = document.add_paragraph()
            p.add_run(f"Soru {i}: ").bold = True
            p.add_run(soru.soru_metni.replace('\r\n', '\n'))
            
            if soru.soru_resmi:
                try:
                    resim_yolu = soru.soru_resmi.path
                    if os.path.exists(resim_yolu):
                        document.add_picture(resim_yolu, width=Inches(2.5))
                except Exception:
                    pass 

            # Sadece içi dolu olan şıkları Word'e ekle
            if soru.secenek_a and str(soru.secenek_a).strip(): document.add_paragraph(f"A) {soru.secenek_a}")
            if soru.secenek_b and str(soru.secenek_b).strip(): document.add_paragraph(f"B) {soru.secenek_b}")
            if soru.secenek_c and str(soru.secenek_c).strip(): document.add_paragraph(f"C) {soru.secenek_c}")
            if soru.secenek_d and str(soru.secenek_d).strip(): document.add_paragraph(f"D) {soru.secenek_d}")
            if soru.secenek_e and str(soru.secenek_e).strip(): document.add_paragraph(f"E) {soru.secenek_e}")
            
            # Eğer hiçbir şık yoksa (Klasik Soru), öğrencinin yazması için Word'de boşluk bırak
            if not any([soru.secenek_a, soru.secenek_b, soru.secenek_c, soru.secenek_d]):
                document.add_paragraph("\n\n\n")
            
            document.add_paragraph("-" * 40)
        
        document.add_page_break() 
        
        cevap_baslik = document.add_paragraph()
        baslik_run = cevap_baslik.add_run('CEVAP ANAHTARI')
        baslik_run.bold = True
        
        # Word için Cevap Anahtarı (Resim ve Klasik Metin Destekli Alt Alta Sıralama)
        for i, soru in enumerate(sorular, 1):
            cevap_p = document.add_paragraph()
            cevap_metni = soru.dogru_cevap if soru.dogru_cevap else "Klasik / Açık Uçlu"
            cevap_p.add_run(f"{i}. Soru: {cevap_metni}").bold = True
            
            # Eğer sorunun klasik referans cevabı varsa Word'e yaz
            if soru.klasik_cevap:
                document.add_paragraph(f"Referans Cevap: {soru.klasik_cevap}")

            # Eğer cevap resmi varsa Word'e ekle
            if soru.cevap_resmi:
                try:
                    resim_yolu = soru.cevap_resmi.path
                    if os.path.exists(resim_yolu):
                        document.add_picture(resim_yolu, width=Inches(2.0))
                except Exception:
                    pass
            
            # Sorular arası estetik boşluk
            document.add_paragraph("")

    f = io.BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename="Matbaa_Akademi_Sinav_Kagidi.docx"'
    response['Content-Length'] = length
    return response


# --- HATA BİLDİR FONKSİYONU ---
def hata_bildir(request):
    """Soru kartındaki hata bildirimlerini kaydeder."""
    if request.method == 'POST':
        soru_id = request.POST.get('soru_id')
        mesaj = request.POST.get('mesaj')

        if not soru_id or not mesaj:
            return JsonResponse({'success': False, 'error': 'Lütfen mesaj alanını doldurun.'})

        soru = get_object_or_404(Soru, id=soru_id)
        kullanici = request.user if request.user.is_authenticated else None

        HataBildirimi.objects.create(soru=soru, kullanici=kullanici, mesaj=mesaj)
        return JsonResponse({'success': True})
        
    return JsonResponse({'success': False, 'error': 'Geçersiz istek.'})